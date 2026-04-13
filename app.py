# backend/app.py
from flask import Flask, request, jsonify
from flask_cors import CORS
import joblib
import os
import re as _re
import requests as http_requests
from dotenv import load_dotenv
import google.generativeai as genai

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
CORS(app)

# Configure Gemini AI
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    gemini_model = genai.GenerativeModel("gemini-2.0-flash")
    print("Gemini AI configured successfully")
else:
    gemini_model = None
    print("WARNING: GEMINI_API_KEY not set. Chatbot will use fallback responses.")

# Paths to saved model and encoder
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MODEL_DIR = os.path.join(BASE_DIR, "models")
model_path = os.path.join(MODEL_DIR, 'career_guidance_pipeline.joblib')
encoder_path = os.path.join(MODEL_DIR, 'label_encoder.joblib')

# Try to load real model/encoder; otherwise use mocks so server still runs
if os.path.exists(model_path) and os.path.exists(encoder_path):
    model = joblib.load(model_path)
    le_interest = joblib.load(encoder_path)
    print(f"Loaded model from {model_path} and encoder from {encoder_path}")
else:
    print("Model files not found! Using mock model and encoder. Run train_model.py to create real models.")
    class MockModel:
        def predict(self, X): 
            # Always returns a sample career path
            return ["Software Engineer"]
    class MockEncoder:
        def transform(self, X): 
            # Returns dummy encoding (0)
            return [0]
    model = MockModel()
    le_interest = MockEncoder()

# Note: Removed unused mappings and ML/NLP imports to keep the backend minimal

# Comprehensive career profiles with courses, roadmaps, and detailed information
career_profiles = {
    "Data Scientist": {
        "companies": [
            {"name": "Google AI", "website": "https://careers.google.com"},
            {"name": "Microsoft Research", "website": "https://careers.microsoft.com"},
            {"name": "Amazon ML", "website": "https://www.amazon.jobs"},
            {"name": "TCS Innovation Labs", "website": "https://www.tcs.com/careers"},
            {"name": "Infosys", "website": "https://www.infosys.com/careers"}
        ],
        "salary": "₹10–18 LPA",
        "skills": ["Python", "Machine Learning", "Statistics", "SQL", "TensorFlow", "Pandas"],
        "courses": [
            {"title": "Machine Learning Specialization", "platform": "Coursera", "instructor": "Andrew Ng"},
            {"title": "Deep Learning A-Z", "platform": "Udemy", "instructor": "Kirill Eremenko"},
            {"title": "Python for Data Science", "platform": "DataCamp", "instructor": "Various"}
        ],
        "roadmap": [
            "Learn Python fundamentals and data manipulation",
            "Master statistics and probability",
            "Complete 3-5 Kaggle projects",
            "Learn TensorFlow and PyTorch",
            "Apply for ML internships",
            "Build a strong portfolio",
            "Apply for Data Scientist roles"
        ]
    },
    "Software Engineer": {
        "companies": [
            {"name": "Google", "website": "https://careers.google.com"},
            {"name": "Microsoft", "website": "https://careers.microsoft.com"},
            {"name": "Amazon", "website": "https://www.amazon.jobs"},
            {"name": "Infosys", "website": "https://www.infosys.com/careers"},
            {"name": "TCS", "website": "https://www.tcs.com/careers"}
        ],
        "salary": "₹6–12 LPA",
        "skills": ["Java", "C++", "OOPs", "Data Structures", "Algorithms", "System Design"],
        "courses": [
            {"title": "CS50: Introduction to Computer Science", "platform": "edX", "instructor": "David J. Malan"},
            {"title": "System Design Primer", "platform": "GitHub", "instructor": "Community"},
            {"title": "Java Programming Masterclass", "platform": "Udemy", "instructor": "Tim Buchalka"}
        ],
        "roadmap": [
            "Master Data Structures and Algorithms",
            "Learn a programming language deeply",
            "Build full-stack projects",
            "Contribute to open-source projects",
            "Learn system design principles",
            "Practice coding interviews",
            "Apply for SDE roles"
        ]
    },
    "Web Developer": {
        "companies": [
            {"name": "Cognizant", "website": "https://careers.cognizant.com"},
            {"name": "Adobe", "website": "https://www.adobe.com/careers.html"},
            {"name": "Zoho", "website": "https://www.zoho.com/careers.html"},
            {"name": "Freshworks", "website": "https://www.freshworks.com/careers"},
            {"name": "Wipro", "website": "https://careers.wipro.com"}
        ],
        "salary": "₹4–9 LPA",
        "skills": ["HTML", "CSS", "JavaScript", "React", "Node.js", "MongoDB"],
        "courses": [
            {"title": "The Complete Web Developer Bootcamp", "platform": "Udemy", "instructor": "Colt Steele"},
            {"title": "React - The Complete Guide", "platform": "Udemy", "instructor": "Maximilian Schwarzmüller"},
            {"title": "Full Stack Web Development", "platform": "Coursera", "instructor": "JHU"}
        ],
        "roadmap": [
            "Learn HTML, CSS, and JavaScript fundamentals",
            "Master a frontend framework (React/Vue/Angular)",
            "Learn backend development (Node.js/Python)",
            "Understand databases (SQL/NoSQL)",
            "Build full-stack projects",
            "Learn deployment and DevOps basics",
            "Apply for web developer positions"
        ]
    },
    "ML Engineer": {
        "companies": [
            {"name": "NVIDIA", "website": "https://www.nvidia.com/en-in/about-nvidia/careers"},
            {"name": "Amazon", "website": "https://www.amazon.jobs"},
            {"name": "OpenAI", "website": "https://openai.com/careers"},
            {"name": "Tesla", "website": "https://www.tesla.com/careers"},
            {"name": "Uber", "website": "https://www.uber.com/careers"}
        ],
        "salary": "₹12–20 LPA",
        "skills": ["Python", "TensorFlow", "Deep Learning", "Data Analysis", "MLOps", "Docker"],
        "courses": [
            {"title": "Deep Learning Specialization", "platform": "Coursera", "instructor": "Andrew Ng"},
            {"title": "Machine Learning Engineering for Production", "platform": "Coursera", "instructor": "Andrew Ng"},
            {"title": "MLOps Fundamentals", "platform": "Coursera", "instructor": "Google Cloud"}
        ],
        "roadmap": [
            "Master machine learning fundamentals",
            "Learn deep learning frameworks (TensorFlow/PyTorch)",
            "Understand MLOps and model deployment",
            "Learn cloud platforms (AWS/GCP/Azure)",
            "Build end-to-end ML projects",
            "Learn containerization (Docker/Kubernetes)",
            "Apply for ML Engineer positions"
        ]
    },
    "Frontend Developer": {
        "companies": [
            {"name": "Adobe", "website": "https://www.adobe.com/careers.html"},
            {"name": "Figma", "website": "https://www.figma.com/careers"},
            {"name": "Canva", "website": "https://www.canva.com/careers"},
            {"name": "Shopify", "website": "https://www.shopify.com/careers"},
            {"name": "Stripe", "website": "https://stripe.com/jobs"}
        ],
        "salary": "₹5–10 LPA",
        "skills": ["HTML", "CSS", "JavaScript", "React", "Vue.js", "TypeScript"],
        "courses": [
            {"title": "Advanced React and Redux", "platform": "Udemy", "instructor": "Stephen Grider"},
            {"title": "JavaScript: The Complete Guide", "platform": "Udemy", "instructor": "Maximilian Schwarzmüller"},
            {"title": "Frontend Web Development", "platform": "Coursera", "instructor": "Meta"}
        ],
        "roadmap": [
            "Master HTML, CSS, and JavaScript",
            "Learn a modern frontend framework",
            "Understand state management",
            "Learn testing frameworks",
            "Master responsive design",
            "Learn performance optimization",
            "Apply for frontend developer roles"
        ]
    },
    "Data Analyst": {
        "companies": [
            {"name": "Microsoft", "website": "https://careers.microsoft.com"},
            {"name": "Amazon", "website": "https://www.amazon.jobs"},
            {"name": "Flipkart", "website": "https://www.flipkartcareers.com"},
            {"name": "Swiggy", "website": "https://careers.swiggy.com"},
            {"name": "Zomato", "website": "https://www.zomato.com/careers"}
        ],
        "salary": "₹6–12 LPA",
        "skills": ["Python", "SQL", "Excel", "Tableau", "Power BI", "Statistics"],
        "courses": [
            {"title": "Google Data Analytics Certificate", "platform": "Coursera", "instructor": "Google"},
            {"title": "Tableau Desktop Specialist", "platform": "Udemy", "instructor": "Kyle Pew"},
            {"title": "SQL for Data Analysis", "platform": "Udemy", "instructor": "Jose Portilla"}
        ],
        "roadmap": [
            "Learn SQL and database fundamentals",
            "Master Excel and data manipulation",
            "Learn data visualization tools",
            "Understand statistical concepts",
            "Learn Python for data analysis",
            "Build analytical projects",
            "Apply for data analyst positions"
        ]
    },
    "AI Researcher": {
        "companies": [
            {"name": "OpenAI", "website": "https://openai.com/careers"},
            {"name": "DeepMind", "website": "https://deepmind.com/careers"},
            {"name": "Google Research", "website": "https://research.google/careers"},
            {"name": "Facebook AI Research", "website": "https://ai.facebook.com/careers"},
            {"name": "Microsoft Research", "website": "https://careers.microsoft.com"}
        ],
        "salary": "₹15–25 LPA",
        "skills": ["Python", "TensorFlow", "PyTorch", "Research", "Mathematics", "Statistics"],
        "courses": [
            {"title": "Deep Learning Specialization", "platform": "Coursera", "instructor": "Andrew Ng"},
            {"title": "CS229: Machine Learning", "platform": "Stanford", "instructor": "Andrew Ng"},
            {"title": "Advanced Deep Learning", "platform": "Coursera", "instructor": "DeepLearning.AI"}
        ],
        "roadmap": [
            "Master advanced mathematics and statistics",
            "Learn deep learning frameworks",
            "Read research papers regularly",
            "Contribute to open-source AI projects",
            "Publish research work",
            "Attend AI conferences and workshops",
            "Apply for research positions"
        ]
    },
    "Cybersecurity Analyst": {
        "companies": [
            {"name": "CrowdStrike", "website": "https://www.crowdstrike.com/careers"},
            {"name": "Palo Alto Networks", "website": "https://www.paloaltonetworks.com/careers"},
            {"name": "Cisco", "website": "https://www.cisco.com/c/en/us/about/careers"},
            {"name": "Wipro", "website": "https://careers.wipro.com"},
            {"name": "HCL", "website": "https://www.hcl.com/careers"}
        ],
        "salary": "₹8–16 LPA",
        "skills": ["Network Security", "Ethical Hacking", "SIEM", "Risk Assessment", "Python", "Linux"],
        "courses": [
            {"title": "Cybersecurity Specialization", "platform": "Coursera", "instructor": "University of Maryland"},
            {"title": "Practical Ethical Hacking", "platform": "Udemy", "instructor": "Heath Adams"},
            {"title": "Network Security", "platform": "edX", "instructor": "Various"}
        ],
        "roadmap": [
            "Learn networking fundamentals",
            "Study cybersecurity basics",
            "Practice on platforms like TryHackMe",
            "Get Security+ certification",
            "Learn incident response",
            "Build security projects",
            "Apply for SOC analyst roles"
        ]
    },
    "Business Analyst": {
        "companies": [
            {"name": "McKinsey", "website": "https://www.mckinsey.com/careers"},
            {"name": "BCG", "website": "https://www.bcg.com/careers"},
            {"name": "Deloitte", "website": "https://www2.deloitte.com/careers"},
            {"name": "Accenture", "website": "https://www.accenture.com/in-en/careers"},
            {"name": "EY", "website": "https://www.ey.com/en_in/careers"}
        ],
        "salary": "₹7–14 LPA",
        "skills": ["Business Analysis", "Data Analysis", "SQL", "Excel", "Power BI", "Communication"],
        "courses": [
            {"title": "Business Analysis Fundamentals", "platform": "Coursera", "instructor": "University of Virginia"},
            {"title": "SQL for Business Intelligence", "platform": "Udemy", "instructor": "Jose Portilla"},
            {"title": "Power BI Desktop", "platform": "Udemy", "instructor": "Maven Analytics"}
        ],
        "roadmap": [
            "Learn business analysis fundamentals",
            "Master data analysis tools",
            "Understand business processes",
            "Learn stakeholder management",
            "Build analytical projects",
            "Get business analysis certification",
            "Apply for business analyst roles"
        ]
    }
}

@app.route('/')
def home():
    return "AI Career Guidance Backend Running"

@app.route('/careers', methods=['GET'])
def get_careers():
    """Get all available career paths with basic info"""
    try:
        careers = []
        for career_name, career_info in career_profiles.items():
            careers.append({
                'name': career_name,
                'salary': career_info['salary'],
                'skills': career_info['skills'][:3],  # Top 3 skills
                'companies_count': len(career_info['companies'])
            })
        
        return jsonify({
            'careers': careers,
            'total': len(careers)
        })
    except Exception as e:
        return jsonify({'error': f'Failed to get careers: {str(e)}'}), 500

@app.route('/career/<career_name>', methods=['GET'])
def get_career_details(career_name):
    """Get detailed information about a specific career"""
    try:
        if career_name not in career_profiles:
            return jsonify({'error': 'Career not found'}), 404
        
        career_info = career_profiles[career_name]
        return jsonify({
            'career': career_name,
            'info': career_info
        })
    except Exception as e:
        return jsonify({'error': f'Failed to get career details: {str(e)}'}), 500

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    try:
        return jsonify({
            'status': 'healthy',
            'model_loaded': os.path.exists(model_path),
            'encoder_loaded': os.path.exists(encoder_path),
            'careers_available': len(career_profiles)
        })
    except Exception as e:
        return jsonify({'status': 'unhealthy', 'error': str(e)}), 500

def calculate_skill_match_score(user_skills, career_skills):
    """Calculate match score based on skills overlap"""
    if not user_skills or not career_skills:
        return 0
    
    user_skills_lower = [skill.lower() for skill in user_skills]
    career_skills_lower = [skill.lower() for skill in career_skills]
    
    matches = sum(1 for skill in user_skills_lower if skill in career_skills_lower)
    return (matches / len(career_skills_lower)) * 100

def calculate_interest_match_score(user_interests, career_name):
    """Calculate match score based on interests and career alignment"""
    if not user_interests:
        return 50  # Default score if no interests provided
    
    # Map interests to career categories
    interest_to_career_mapping = {
        'Artificial Intelligence': ['Data Scientist', 'ML Engineer', 'AI Researcher'],
        'Web Development': ['Web Developer', 'Frontend Developer', 'Software Engineer'],
        'Mobile Apps': ['Mobile Developer', 'Software Engineer'],
        'Data Science': ['Data Scientist', 'Data Analyst', 'ML Engineer'],
        'Cybersecurity': ['Cybersecurity Analyst'],
        'Entrepreneurship': ['Business Analyst', 'Software Engineer'],
        'Research': ['AI Researcher', 'Research Scientist', 'Data Scientist'],
        'Business': ['Business Analyst', 'Data Analyst']
    }
    
    # Check if any user interest aligns with the career
    for interest in user_interests:
        if interest in interest_to_career_mapping:
            if career_name in interest_to_career_mapping[interest]:
                return 90  # High match
    return 60  # Medium match

def normalize_cgpa(cgpa):
    """Normalize CGPA to 0-1 scale"""
    try:
        cgpa_val = float(cgpa)
        if cgpa_val <= 10:
            return cgpa_val / 10
        else:
            return min(cgpa_val / 100, 1)
    except (ValueError, TypeError):
        return 0.5  # Default normalized score

def get_career_recommendations(user_data):
    """Get comprehensive career recommendations based on user profile"""
    user_skills = user_data.get('skills', [])
    user_interests = user_data.get('interests', [])
    cgpa = user_data.get('cgpa', 0)
    
    # Normalize CGPA
    cgpa_normalized = normalize_cgpa(cgpa)
    
    # Calculate scores for all careers
    career_scores = []
    
    for career_name, career_info in career_profiles.items():
        # Calculate skill match (40% weight)
        skill_score = calculate_skill_match_score(user_skills, career_info['skills'])
        
        # Calculate interest match (30% weight)
        interest_score = calculate_interest_match_score(user_interests, career_name)
        
        # Calculate CGPA score (30% weight)
        cgpa_score = cgpa_normalized * 100
        
        # Weighted total score
        total_score = (skill_score * 0.4) + (interest_score * 0.3) + (cgpa_score * 0.3)
        
        career_scores.append({
            'career': career_name,
            'score': round(total_score),
            'skill_match': round(skill_score),
            'interest_match': round(interest_score),
            'cgpa_score': round(cgpa_score),
            'info': career_info
        })
    
    # Sort by score (highest first)
    career_scores.sort(key=lambda x: x['score'], reverse=True)
    
    return career_scores

@app.route('/predict', methods=['POST'])
def predict():
    """
    Enhanced prediction endpoint that handles multiple skills and interests
    
    Expects JSON payload:
    {
      "name": "Alice",
      "cgpa": 8.3,
      "skills": ["Python", "Machine Learning"],
      "interests": ["Artificial Intelligence", "Data Science"]
    }

    Returns:
    {
      "career": "Data Scientist",
      "info": { "companies": [...], "salary": "...", "skills": [...], "courses": [...], "roadmap": [...] },
      "alternatives": [...],
      "match_percentage": 85
    }
    """
    data = request.get_json(force=True, silent=True)
    if not data:
        return jsonify({'error': 'Invalid or missing JSON payload'}), 400

    # Mapping skills to likely interests
    SKILL_TO_INTERESTS = {
        'Python': ['Data Science', 'Artificial Intelligence', 'Web Development'],
        'Java': ['Web Development', 'Mobile Apps', 'Entrepreneurship'],
        'JavaScript': ['Web Development', 'Mobile Apps'],
        'React': ['Web Development', 'Mobile Apps'],
        'Machine Learning': ['Artificial Intelligence', 'Data Science', 'Research'],
        'Data Analysis': ['Data Science', 'Business', 'Research'],
        'IoT': ['Research', 'Entrepreneurship'],
        'Cloud Computing': ['Web Development', 'Entrepreneurship'],
        'UI/UX Design': ['Web Development', 'Entrepreneurship'],
        'Digital Marketing': ['Business', 'Entrepreneurship'],
    }

    try:
        # Extract user data
        name = data.get('name', 'Anonymous')
        cgpa = float(data.get('cgpa', 0))
        skills = data.get('skills', [])
        interests = data.get('interests', [])
        
        # Auto-derive interests from skills when none provided
        if not interests:
            derived = []
            for skill in skills:
                derived.extend(SKILL_TO_INTERESTS.get(skill, []))
            interests = list(dict.fromkeys(derived))  # dedupe, preserve order
        
        # Validate input
        if not skills:
            return jsonify({'error': 'Please provide at least one skill'}), 400
        
        # Get career recommendations
        recommendations = get_career_recommendations({
            'skills': skills,
            'interests': interests,
            'cgpa': cgpa
        })
        
        if not recommendations:
            return jsonify({'error': 'No career recommendations found'}), 400
        
        # Primary recommendation
        primary = recommendations[0]
        
        # Alternative recommendations (top 2-3 excluding primary)
        alternatives = recommendations[1:4]  # Get next 3 alternatives
        
        # Prepare response
        response = {
            'career': primary['career'],
            'match_percentage': primary['score'],
            'info': {
                'companies': primary['info']['companies'],
                'salary': primary['info']['salary'],
                'skills': primary['info']['skills'],
                'courses': primary['info']['courses'],
                'roadmap': primary['info']['roadmap']
            },
            'suggested_interests': interests,
            'alternatives': [
                {
                    'career': alt['career'],
                    'match_percentage': alt['score']
                } for alt in alternatives
            ],
            'analysis': {
                'skill_match': primary['skill_match'],
                'interest_match': primary['interest_match'],
                'cgpa_score': primary['cgpa_score']
            }
        }
        
        return jsonify(response)
        
    except ValueError as e:
        return jsonify({'error': f'Invalid data format: {str(e)}'}), 400
    except Exception as e:
        return jsonify({'error': f'Prediction failed: {str(e)}'}), 500

# ---------- Resume Analysis ----------

# Skills dictionary for resume parsing
RESUME_SKILLS_DICT = [
    # Programming Languages
    "Python", "Java", "JavaScript", "TypeScript", "C", "C++", "C#", "Go", "Rust", "Ruby",
    "PHP", "Swift", "Kotlin", "Scala", "R", "MATLAB", "Perl", "Shell", "Bash",
    # Web Development
    "HTML", "CSS", "React", "Angular", "Vue", "Vue.js", "Node.js", "Express", "Django",
    "Flask", "Spring", "Spring Boot", "Next.js", "Nuxt.js", "Svelte", "jQuery",
    "Bootstrap", "Tailwind", "SASS", "LESS", "WordPress",
    # Data & ML
    "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Keras", "Scikit-learn",
    "Pandas", "NumPy", "Data Analysis", "Data Science", "NLP", "Computer Vision",
    "Reinforcement Learning", "Neural Networks", "Random Forest", "SVM",
    "Natural Language Processing", "Big Data", "Hadoop", "Spark", "Tableau", "Power BI",
    "Excel", "Statistics", "Data Visualization", "Data Mining", "Data Engineering",
    # Databases
    "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Firebase", "Cassandra",
    "DynamoDB", "Oracle", "SQLite", "NoSQL", "GraphQL",
    # Cloud & DevOps
    "AWS", "Azure", "GCP", "Google Cloud", "Docker", "Kubernetes", "CI/CD",
    "Jenkins", "Terraform", "Ansible", "Linux", "Git", "GitHub", "GitLab",
    "Cloud Computing", "Microservices", "REST API", "API Development",
    # Mobile
    "Android", "iOS", "React Native", "Flutter", "SwiftUI", "Xamarin",
    # Security
    "Cybersecurity", "Network Security", "Ethical Hacking", "Penetration Testing",
    "SIEM", "Risk Assessment", "Cryptography", "OWASP",
    # Design
    "UI/UX Design", "Figma", "Adobe XD", "Sketch", "Photoshop", "Illustrator",
    "User Research", "Wireframing", "Prototyping",
    # Business & Marketing
    "Digital Marketing", "SEO", "SEM", "Google Analytics", "Content Marketing",
    "Social Media Marketing", "Project Management", "Agile", "Scrum", "JIRA",
    "Business Analysis", "Communication", "Leadership",
    # IoT & Embedded
    "IoT", "Arduino", "Raspberry Pi", "Embedded Systems",
    # Other
    "Blockchain", "Smart Contracts", "Solidity", "Game Development", "Unity", "Unreal Engine",
    "3D Modeling", "Blender", "AutoCAD", "SolidWorks",
]

def extract_text_from_pdf(file_stream):
    """Extract text from a PDF file stream"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_stream)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {str(e)}")

def extract_text_from_docx(file_stream):
    """Extract text from a DOCX file stream"""
    try:
        from docx import Document
        doc = Document(file_stream)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {str(e)}")

def extract_skills_from_text(text):
    """Extract skills from resume text using keyword matching"""
    text_lower = text.lower()
    found_skills = []
    for skill in RESUME_SKILLS_DICT:
        # Use word boundary matching for shorter skill names
        skill_lower = skill.lower()
        if len(skill_lower) <= 3:
            # For short skills like "R", "C", "Go", "SQL" use word boundary regex
            pattern = r'\b' + _re.escape(skill_lower) + r'\b'
            if _re.search(pattern, text_lower):
                found_skills.append(skill)
        else:
            if skill_lower in text_lower:
                found_skills.append(skill)
    return list(dict.fromkeys(found_skills))  # dedupe, preserve order

def extract_cgpa_from_text(text):
    """Extract CGPA/GPA from resume text using regex patterns"""
    patterns = [
        r'(?:CGPA|GPA|C\.G\.P\.A|G\.P\.A)\s*[:\-]?\s*(\d+\.?\d*)\s*(?:/\s*(\d+))?',
        r'(?:CGPA|GPA)\s*(\d+\.?\d*)',
        r'(\d+\.?\d*)\s*/\s*(?:10|4)\s*(?:CGPA|GPA)',
        r'(?:percentage|percent|%)\s*[:\-]?\s*(\d+\.?\d*)',
        r'(\d+\.?\d*)\s*%',
    ]
    
    for pattern in patterns:
        match = _re.search(pattern, text, _re.IGNORECASE)
        if match:
            value = float(match.group(1))
            # If it has a scale denominator
            if match.lastindex and match.lastindex >= 2 and match.group(2):
                scale = float(match.group(2))
                if scale == 4:
                    value = value * 2.5  # Convert 4.0 scale to 10.0
            # Validate range
            if 0 < value <= 10:
                return value
            elif 10 < value <= 100:
                return value  # Percentage
    return 7.0  # Default CGPA if not found

def extract_name_from_text(text):
    """Try to extract candidate name from the first few lines of resume"""
    lines = text.strip().split('\n')
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if not line:
            continue
        # Skip lines that look like headers, emails, phones, URLs
        if _re.search(r'@|http|www\.|resume|curriculum|vitae|\d{5,}', line, _re.IGNORECASE):
            continue
        # Skip lines that are too long (likely paragraphs)
        if len(line) > 50:
            continue
        # Check if it looks like a name (2-4 words, all alpha)
        words = line.split()
        if 1 <= len(words) <= 4 and all(_re.match(r'^[a-zA-Z\.\-]+$', w) for w in words):
            return line
    return "Resume Candidate"


@app.route('/analyze-resume', methods=['POST'])
def analyze_resume():
    """
    Analyze an uploaded resume file (PDF or DOCX).
    
    Expects multipart/form-data with a 'resume' file field.
    
    Returns the same JSON structure as /predict.
    """
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file provided'}), 400
    
    file = request.files['resume']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
    
    filename = file.filename.lower()
    
    try:
        # Extract text based on file type
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file.stream)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file.stream)
        elif filename.endswith('.txt'):
            text = file.stream.read().decode('utf-8', errors='ignore')
        else:
            return jsonify({'error': 'Unsupported file type. Please upload PDF, DOCX, or TXT.'}), 400
        
        if not text or len(text.strip()) < 20:
            return jsonify({'error': 'Could not extract enough text from the file. Please try a different file.'}), 400
        
        # Extract information from resume
        skills = extract_skills_from_text(text)
        cgpa = extract_cgpa_from_text(text)
        name = extract_name_from_text(text)
        
        if not skills:
            return jsonify({'error': 'No recognizable skills found in the resume. Please ensure your resume lists your technical skills.'}), 400
        
        # Derive interests from skills
        SKILL_TO_INTERESTS = {
            'Python': ['Data Science', 'Artificial Intelligence', 'Web Development'],
            'Java': ['Web Development', 'Mobile Apps', 'Entrepreneurship'],
            'JavaScript': ['Web Development', 'Mobile Apps'],
            'React': ['Web Development', 'Mobile Apps'],
            'Machine Learning': ['Artificial Intelligence', 'Data Science', 'Research'],
            'Data Analysis': ['Data Science', 'Business', 'Research'],
            'IoT': ['Research', 'Entrepreneurship'],
            'Cloud Computing': ['Web Development', 'Entrepreneurship'],
            'UI/UX Design': ['Web Development', 'Entrepreneurship'],
            'Digital Marketing': ['Business', 'Entrepreneurship'],
            'Deep Learning': ['Artificial Intelligence', 'Data Science', 'Research'],
            'TensorFlow': ['Artificial Intelligence', 'Data Science'],
            'PyTorch': ['Artificial Intelligence', 'Data Science'],
            'Docker': ['Web Development', 'Entrepreneurship'],
            'Kubernetes': ['Web Development', 'Entrepreneurship'],
            'AWS': ['Web Development', 'Entrepreneurship'],
            'Azure': ['Web Development', 'Entrepreneurship'],
            'Cybersecurity': ['Cybersecurity'],
            'Network Security': ['Cybersecurity'],
            'Flutter': ['Mobile Apps'],
            'React Native': ['Mobile Apps'],
            'SQL': ['Data Science', 'Business'],
            'Tableau': ['Data Science', 'Business'],
            'Power BI': ['Data Science', 'Business'],
        }
        
        interests = []
        for skill in skills:
            interests.extend(SKILL_TO_INTERESTS.get(skill, []))
        interests = list(dict.fromkeys(interests))
        
        # Get career recommendations
        recommendations = get_career_recommendations({
            'skills': skills,
            'interests': interests,
            'cgpa': cgpa
        })
        
        if not recommendations:
            return jsonify({'error': 'No career recommendations found'}), 400
        
        primary = recommendations[0]
        alternatives = recommendations[1:4]
        
        response = {
            'career': primary['career'],
            'match_percentage': primary['score'],
            'info': {
                'companies': primary['info']['companies'],
                'salary': primary['info']['salary'],
                'skills': primary['info']['skills'],
                'courses': primary['info']['courses'],
                'roadmap': primary['info']['roadmap']
            },
            'suggested_interests': interests,
            'alternatives': [
                {
                    'career': alt['career'],
                    'match_percentage': alt['score']
                } for alt in alternatives
            ],
            'analysis': {
                'skill_match': primary['skill_match'],
                'interest_match': primary['interest_match'],
                'cgpa_score': primary['cgpa_score']
            },
            'extracted_data': {
                'name': name,
                'skills_found': skills,
                'cgpa_detected': cgpa,
                'resume_text_length': len(text)
            }
        }
        
        return jsonify(response)
        
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f'Resume analysis failed: {str(e)}'}), 500


# ---------- Roadmap Integration ----------



CAREER_TO_ROADMAP_SLUG = {
    "Data Scientist": "ai-data-scientist",
    "Software Engineer": "backend",
    "Web Developer": "frontend",
    "ML Engineer": "mlops",
    "Frontend Developer": "frontend",
    "Data Analyst": "data-analyst",
    "AI Researcher": "ai-data-scientist",
    "Cybersecurity Analyst": "cyber-security",
    "Business Analyst": "data-analyst",
    "UI/UX Designer": "ux-design",
    "Cloud Engineer": "devops",
}

GITHUB_RAW_BASE = "https://raw.githubusercontent.com/kamranahmedse/developer-roadmap/master/src/data/roadmaps"

def _label_to_slug(label):
    """Convert a node label to a content file slug"""
    slug = label.lower().strip()
    slug = _re.sub(r'[^a-z0-9\s-]', '', slug)
    slug = _re.sub(r'\s+', '-', slug)
    return slug

def _parse_content_md(md_text):
    """Parse roadmap.sh content markdown into structured data"""
    lines = md_text.strip().split('\n')
    title = ''
    description_lines = []
    resources = []
    in_resources = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('# '):
            title = stripped[2:].strip()
            continue
        if 'visit the following resources' in stripped.lower():
            in_resources = True
            continue
        if in_resources and stripped.startswith('- ['):
            # Parse resource line: - [@type@Title](URL)
            match = _re.match(r'-\s*\[@(\w+)@(.+?)\]\((.+?)\)', stripped)
            if match:
                rtype, rtitle, rurl = match.groups()
                resources.append({
                    'type': rtype,
                    'title': rtitle,
                    'url': rurl
                })
        elif not in_resources and stripped:
            description_lines.append(stripped)

    return {
        'title': title,
        'description': ' '.join(description_lines),
        'resources': resources
    }


@app.route('/roadmap/<career_name>', methods=['GET'])
def get_roadmap(career_name):
    """Fetch full roadmap node graph for visual rendering"""
    try:
        slug = CAREER_TO_ROADMAP_SLUG.get(career_name)
        if not slug:
            return jsonify({'error': f'No roadmap mapping for "{career_name}"'}), 404

        url = f"https://roadmap.sh/{slug}.json"
        resp = http_requests.get(url, timeout=15)
        if resp.status_code != 200:
            return jsonify({'error': 'Failed to fetch roadmap data'}), 502

        data = resp.json()
        raw_nodes = data.get('nodes', [])
        raw_edges = data.get('edges', [])
        title_obj = data.get('title', {})
        roadmap_title = title_obj.get('page', title_obj.get('card', career_name))

        # Build nodes for the frontend
        nodes = []
        for node in raw_nodes:
            ntype = node.get('type', '')
            label = (node.get('data') or {}).get('label', '')
            if not label or ntype not in ('topic', 'subtopic', 'title'):
                continue

            pos = node.get('position', {})
            style = node.get('style', {})
            data_style = (node.get('data') or {}).get('style', {})
            legend = (node.get('data') or {}).get('legend', {})

            nodes.append({
                'id': node.get('id', ''),
                'type': ntype,
                'label': label,
                'x': pos.get('x', 0),
                'y': pos.get('y', 0),
                'width': style.get('width', node.get('width', 240)),
                'height': style.get('height', node.get('height', 49)),
                'legendColor': legend.get('color', ''),
                'legendLabel': legend.get('label', ''),
            })

        # Sort by y then x
        nodes.sort(key=lambda n: (n['y'], n['x']))

        # Build edges
        edges = []
        for edge in raw_edges:
            edge_style = edge.get('style', {})
            edges.append({
                'source': edge.get('source', ''),
                'target': edge.get('target', ''),
                'dashed': '8' in str(edge_style.get('strokeDasharray', '0')),
            })

        # Build sections (grouped topics with subtopics) for a simpler view
        topics = []
        subtopics_list = []
        for n in nodes:
            if n['type'] == 'topic':
                topics.append({**n, 'subtopics': []})
            elif n['type'] == 'subtopic':
                subtopics_list.append(n)

        for sub in subtopics_list:
            best = None
            best_d = float('inf')
            for t in topics:
                d = abs(sub['y'] - t['y'])
                if d < best_d:
                    best_d = d
                    best = t
            if best:
                best['subtopics'].append(sub)

        sections = []
        for t in topics:
            sections.append({
                'id': t['id'],
                'topic': t['label'],
                'subtopics': [{'id': s['id'], 'label': s['label']} for s in t['subtopics']]
            })

        return jsonify({
            'title': roadmap_title,
            'slug': slug,
            'sections': sections,
            'nodes': nodes,
            'edges': edges,
        })

    except http_requests.exceptions.Timeout:
        return jsonify({'error': 'Request timed out'}), 504
    except Exception as e:
        return jsonify({'error': f'Failed to get roadmap: {str(e)}'}), 500


@app.route('/roadmap/<career_name>/content/<node_id>', methods=['GET'])
def get_roadmap_content(career_name, node_id):
    """Fetch content/resources for a specific roadmap node from GitHub"""
    try:
        slug = CAREER_TO_ROADMAP_SLUG.get(career_name)
        if not slug:
            return jsonify({'error': 'Career not found'}), 404

        # First, get the roadmap JSON to find the node's label
        url = f"https://roadmap.sh/{slug}.json"
        resp = http_requests.get(url, timeout=10)
        if resp.status_code != 200:
            return jsonify({'error': 'Failed to fetch roadmap'}), 502

        data = resp.json()
        node_label = None
        for node in data.get('nodes', []):
            if node.get('id') == node_id:
                node_label = (node.get('data') or {}).get('label', '')
                break

        if not node_label:
            return jsonify({'error': 'Node not found', 'title': '', 'description': '', 'resources': []}), 404

        # Try to fetch content from GitHub
        content_slug = _label_to_slug(node_label)
        content_url = f"{GITHUB_RAW_BASE}/{slug}/content/{content_slug}@{node_id}.md"
        
        content_resp = http_requests.get(content_url, timeout=10)
        if content_resp.status_code == 200:
            parsed = _parse_content_md(content_resp.text)
            return jsonify(parsed)
        
        # Fallback: return just the label as title
        return jsonify({
            'title': node_label,
            'description': f'Learn about {node_label} to advance in your career path.',
            'resources': []
        })

    except Exception as e:
        return jsonify({
            'title': 'Content unavailable',
            'description': str(e),
            'resources': []
        }), 500


# ---------- AI Chatbot ----------

CHAT_SYSTEM_PROMPT = """You are an expert AI Career Counselor for computer science and IT students in India.
Your role is to provide helpful, concise, and actionable career guidance.

Topics you can help with:
- Career path recommendations (Data Science, Software Engineering, Web Dev, ML, Cybersecurity, etc.)
- Salary expectations in India (provide ranges in LPA)
- Skills to learn and skill development roadmaps
- Interview preparation tips
- Resume building advice
- Course and certification recommendations
- Industry trends and job market insights
- Internship and job search strategies

Guidelines:
- Keep responses concise (2-4 sentences max unless asked for details)
- Use emojis occasionally to keep it friendly
- Always relate advice to practical, actionable steps
- When mentioning salaries, use Indian Rupee (₹) and LPA format
- If asked something unrelated to careers/tech, politely redirect to career topics
"""

# Fallback knowledge base for when Gemini is unavailable
CHAT_KB_FALLBACK = {
    "salary": "💰 Salaries: Data Scientist ₹10-18 LPA, Software Engineer ₹6-12 LPA, ML Engineer ₹12-20 LPA, Cybersecurity ₹8-16 LPA.",
    "skills": "🎯 Top skills: Python, JavaScript, React, ML, Cloud (AWS/Azure), SQL, Docker. Build projects to demonstrate them!",
    "interview": "📝 Tips: 1) Practice DSA on LeetCode 2) Build portfolio 3) STAR-format answers 4) Research company 5) System design prep.",
    "resume": "📄 Resume: 1 page, action verbs, quantify achievements, GitHub links, tailor per job, add certifications.",
    "career": "🚀 Use our AI tool! Identify strengths → explore domains → do internships → network → get AI recommendations.",
    "python": "🐍 Python: Great for Data Science, ML, Web (Django/Flask), Automation. Most versatile language!",
    "machine learning": "🤖 ML: Python+Math → algorithms → Kaggle → TensorFlow/PyTorch → projects → apply. Try Andrew Ng courses!",
    "web": "🌐 Web Dev: HTML/CSS → JS → React → Node.js → DBs → Deploy. Build 3-5 portfolio projects.",
    "cloud": "☁️ Cloud: AWS/GCP/Azure basics → certify → Docker/K8s → Terraform → apply for cloud roles.",
}


@app.route('/chat', methods=['POST'])
def chat():
    """AI Chatbot endpoint powered by Gemini API"""
    data = request.get_json(force=True, silent=True)
    if not data or not data.get('message', '').strip():
        return jsonify({'error': 'Message is required'}), 400

    user_message = data['message'].strip()
    chat_history = data.get('history', [])

    # Try Gemini API first
    if gemini_model:
        try:
            # Build conversation history for context
            contents = []
            # Add system prompt as first user message with model acknowledgment
            contents.append({"role": "user", "parts": [CHAT_SYSTEM_PROMPT + "\n\nPlease acknowledge and follow these instructions."]})
            contents.append({"role": "model", "parts": ["Understood! I'm your AI Career Counselor. I'll provide helpful, concise career guidance for CS/IT students in India. Ask me anything about careers, skills, salaries, interviews, or job search strategies! 🚀"]})

            # Add chat history (last 10 messages for context window)
            for msg in chat_history[-10:]:
                role = "user" if msg.get('from') == 'user' else "model"
                contents.append({"role": role, "parts": [msg.get('text', '')]})

            # Add current message
            contents.append({"role": "user", "parts": [user_message]})

            response = gemini_model.generate_content(contents)
            reply = response.text.strip()

            return jsonify({'reply': reply, 'source': 'gemini'})

        except Exception as e:
            print(f"Gemini API error: {e}")
            # Fall through to fallback

    # Fallback: keyword-based responses
    lo = user_message.lower()
    reply = "🤔 I can help with: salary info, skills advice, interview tips, resume building, career paths, Python, ML, web dev, and cloud computing. Ask me anything!"
    for keyword, response in CHAT_KB_FALLBACK.items():
        if keyword in lo:
            reply = response
            break

    return jsonify({'reply': reply, 'source': 'fallback'})


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_DEBUG", "True").lower() in ("1", "true", "yes")
    app.run(host='0.0.0.0', port=port, debug=debug)
